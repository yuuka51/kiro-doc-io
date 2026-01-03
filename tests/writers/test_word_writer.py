"""WordライターのユニットテストWord

要件6.1: 構造化データからWordファイルを生成する
要件6.2: 見出し、段落、表、箇条書きリストを含むドキュメントを作成する
要件6.3: 見出しレベル（H1、H2、H3など）を適切に適用する
要件6.4: 生成したファイルをユーザーがアクセス可能な場所に保存する
要件6.5: ファイル生成中にエラーが発生した場合、エラーメッセージを返す
"""

import os
import pytest
import tempfile

from docx import Document

from document_format_mcp_server.writers.word_writer import WordWriter
from document_format_mcp_server.utils.models import WriteResult


@pytest.fixture
def word_writer():
    """WordWriterインスタンスを返すフィクスチャ"""
    return WordWriter()


@pytest.fixture
def temp_output_dir():
    """一時出力ディレクトリを返すフィクスチャ"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def basic_document_data():
    """基本的なドキュメントデータを返すフィクスチャ"""
    return {
        "title": "テストドキュメント",
        "sections": [
            {
                "heading": "セクション1",
                "level": 1,
                "paragraphs": ["これはテスト段落です。"]
            }
        ]
    }


@pytest.fixture
def multi_section_data():
    """複数セクションのデータを返すフィクスチャ"""
    return {
        "title": "複数セクションドキュメント",
        "sections": [
            {
                "heading": "概要",
                "level": 1,
                "paragraphs": ["このドキュメントの概要です。"]
            },
            {
                "heading": "詳細",
                "level": 1,
                "paragraphs": ["詳細な説明です。", "追加の段落です。"]
            },
            {
                "heading": "サブセクション",
                "level": 2,
                "paragraphs": ["サブセクションの内容です。"]
            }
        ]
    }


class TestWordWriterBasic:
    """WordWriterの基本機能テスト"""

    def test_create_document_returns_write_result(
        self, word_writer, temp_output_dir, basic_document_data
    ):
        """WordWriterがWriteResultを返すことを検証"""
        output_path = os.path.join(temp_output_dir, "test.docx")
        result = word_writer.create_document(basic_document_data, output_path)
        
        # WriteResultデータクラスであることを検証
        assert isinstance(result, WriteResult)
        assert hasattr(result, 'success')
        assert hasattr(result, 'output_path')
        assert hasattr(result, 'url')
        assert hasattr(result, 'error')

    def test_create_document_success(
        self, word_writer, temp_output_dir, basic_document_data
    ):
        """Wordファイルが正常に生成されることを検証（要件6.1）"""
        output_path = os.path.join(temp_output_dir, "test.docx")
        result = word_writer.create_document(basic_document_data, output_path)
        
        # 成功フラグを検証
        assert result.success is True
        assert result.error is None
        assert result.output_path == output_path
        
        # ファイルが存在することを検証
        assert os.path.exists(output_path)

    def test_create_document_file_is_valid_docx(
        self, word_writer, temp_output_dir, basic_document_data
    ):
        """生成されたファイルが有効なWordファイルであることを検証"""
        output_path = os.path.join(temp_output_dir, "test.docx")
        word_writer.create_document(basic_document_data, output_path)
        
        # python-docxで読み込めることを検証
        doc = Document(output_path)
        assert doc is not None
        # 段落が存在することを検証
        assert len(doc.paragraphs) > 0


class TestWordWriterSections:
    """セクション作成のテスト"""

    def test_create_section_with_heading(self, word_writer, temp_output_dir):
        """見出し付きセクションが正しく作成されることを検証（要件6.2）"""
        data = {
            "sections": [
                {
                    "heading": "テスト見出し",
                    "level": 1,
                    "paragraphs": ["テスト段落"]
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "heading_test.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        # ドキュメントの内容を検証
        doc = Document(output_path)
        # 見出しが含まれていることを検証
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 1

    def test_create_section_with_paragraphs(self, word_writer, temp_output_dir):
        """段落が正しく作成されることを検証（要件6.2）"""
        data = {
            "sections": [
                {
                    "heading": "セクション",
                    "level": 1,
                    "paragraphs": ["段落1", "段落2", "段落3"]
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "paragraphs_test.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        # ドキュメントの内容を検証
        doc = Document(output_path)
        # 段落が含まれていることを検証
        normal_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Normal']
        assert len(normal_paragraphs) >= 3

    def test_create_section_with_bullets(self, word_writer, temp_output_dir):
        """箇条書きリストが正しく作成されることを検証（要件6.2）"""
        data = {
            "sections": [
                {
                    "heading": "箇条書きセクション",
                    "level": 1,
                    "paragraphs": [],
                    "bullets": ["項目1", "項目2", "項目3"]
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "bullets_test.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        # ドキュメントの内容を検証
        doc = Document(output_path)
        # 箇条書きスタイルの段落が含まれていることを検証
        bullet_paragraphs = [p for p in doc.paragraphs if 'List' in p.style.name]
        assert len(bullet_paragraphs) >= 3

    def test_create_section_with_table(self, word_writer, temp_output_dir):
        """表が正しく作成されることを検証（要件6.2）"""
        data = {
            "sections": [
                {
                    "heading": "表セクション",
                    "level": 1,
                    "paragraphs": [],
                    "tables": [
                        {
                            "data": [
                                ["ヘッダー1", "ヘッダー2"],
                                ["データ1", "データ2"],
                                ["データ3", "データ4"]
                            ]
                        }
                    ]
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "table_test.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        # ドキュメントの内容を検証
        doc = Document(output_path)
        # 表が含まれていることを検証
        assert len(doc.tables) >= 1
        # 表のサイズを検証
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 2

    def test_create_multiple_sections(
        self, word_writer, temp_output_dir, multi_section_data
    ):
        """複数セクションが正しく作成されることを検証（要件6.2）"""
        output_path = os.path.join(temp_output_dir, "multi_sections.docx")
        result = word_writer.create_document(multi_section_data, output_path)
        
        assert result.success is True
        
        # ドキュメントの内容を検証
        doc = Document(output_path)
        # 複数の見出しが含まれていることを検証
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 3


class TestWordWriterHeadingLevels:
    """見出しレベルのテスト"""

    def test_heading_level_1(self, word_writer, temp_output_dir):
        """見出しレベル1が正しく適用されることを検証（要件6.3）"""
        data = {
            "sections": [
                {
                    "heading": "レベル1見出し",
                    "level": 1,
                    "paragraphs": []
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "level1.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        doc = Document(output_path)
        headings = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        assert len(headings) >= 1

    def test_heading_level_2(self, word_writer, temp_output_dir):
        """見出しレベル2が正しく適用されることを検証（要件6.3）"""
        data = {
            "sections": [
                {
                    "heading": "レベル2見出し",
                    "level": 2,
                    "paragraphs": []
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "level2.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        doc = Document(output_path)
        headings = [p for p in doc.paragraphs if p.style.name == 'Heading 2']
        assert len(headings) >= 1

    def test_heading_level_3(self, word_writer, temp_output_dir):
        """見出しレベル3が正しく適用されることを検証（要件6.3）"""
        data = {
            "sections": [
                {
                    "heading": "レベル3見出し",
                    "level": 3,
                    "paragraphs": []
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "level3.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        doc = Document(output_path)
        headings = [p for p in doc.paragraphs if p.style.name == 'Heading 3']
        assert len(headings) >= 1

    def test_mixed_heading_levels(self, word_writer, temp_output_dir):
        """複数の見出しレベルが正しく適用されることを検証（要件6.3）"""
        data = {
            "sections": [
                {"heading": "H1見出し", "level": 1, "paragraphs": []},
                {"heading": "H2見出し", "level": 2, "paragraphs": []},
                {"heading": "H3見出し", "level": 3, "paragraphs": []}
            ]
        }
        output_path = os.path.join(temp_output_dir, "mixed_levels.docx")
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        
        doc = Document(output_path)
        h1 = [p for p in doc.paragraphs if p.style.name == 'Heading 1']
        h2 = [p for p in doc.paragraphs if p.style.name == 'Heading 2']
        h3 = [p for p in doc.paragraphs if p.style.name == 'Heading 3']
        assert len(h1) >= 1
        assert len(h2) >= 1
        assert len(h3) >= 1


class TestWordWriterFileSaving:
    """ファイル保存機能のテスト"""

    def test_save_to_specified_path(
        self, word_writer, temp_output_dir, basic_document_data
    ):
        """指定されたパスにファイルが保存されることを検証（要件6.4）"""
        output_path = os.path.join(temp_output_dir, "specified_path.docx")
        result = word_writer.create_document(basic_document_data, output_path)
        
        assert result.success is True
        assert result.output_path == output_path
        assert os.path.exists(output_path)

    def test_create_output_directory_if_not_exists(
        self, word_writer, temp_output_dir, basic_document_data
    ):
        """出力ディレクトリが存在しない場合に作成されることを検証（要件6.4）"""
        nested_dir = os.path.join(temp_output_dir, "nested", "dir")
        output_path = os.path.join(nested_dir, "test.docx")
        
        result = word_writer.create_document(basic_document_data, output_path)
        
        assert result.success is True
        assert os.path.exists(output_path)
        assert os.path.isdir(nested_dir)


class TestWordWriterErrorHandling:
    """エラーハンドリングのテスト"""

    def test_error_on_missing_sections_key(self, word_writer, temp_output_dir):
        """sectionsキーがない場合にエラーを返すことを検証（要件6.5）"""
        data = {"title": "タイトルのみ"}
        output_path = os.path.join(temp_output_dir, "error.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is False
        assert result.error is not None
        assert "sections" in result.error.lower() or "キー" in result.error

    def test_error_on_invalid_data_type(self, word_writer, temp_output_dir):
        """不正なデータ型の場合にエラーを返すことを検証（要件6.5）"""
        output_path = os.path.join(temp_output_dir, "error.docx")
        
        result = word_writer.create_document("invalid_data", output_path)
        
        assert result.success is False
        assert result.error is not None

    def test_error_on_sections_not_list(self, word_writer, temp_output_dir):
        """sectionsがリストでない場合にエラーを返すことを検証（要件6.5）"""
        data = {"sections": "not_a_list"}
        output_path = os.path.join(temp_output_dir, "error.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is False
        assert result.error is not None

    def test_error_on_invalid_heading_level(self, word_writer, temp_output_dir):
        """不正な見出しレベルの場合にエラーを返すことを検証（要件6.5）"""
        data = {
            "sections": [
                {
                    "heading": "テスト",
                    "level": 5,  # 不正なレベル（1-3のみ有効）
                    "paragraphs": []
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "error.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is False
        assert result.error is not None


class TestWordWriterEdgeCases:
    """エッジケースのテスト"""

    def test_empty_sections_list(self, word_writer, temp_output_dir):
        """空のセクションリストでもファイルが生成されることを検証"""
        data = {"sections": []}
        output_path = os.path.join(temp_output_dir, "empty_sections.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        assert os.path.exists(output_path)

    def test_section_with_empty_heading(self, word_writer, temp_output_dir):
        """空の見出しでもセクションが作成されることを検証"""
        data = {
            "sections": [
                {
                    "heading": "",
                    "level": 1,
                    "paragraphs": ["コンテンツのみ"]
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "empty_heading.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        assert os.path.exists(output_path)

    def test_section_with_empty_paragraphs(self, word_writer, temp_output_dir):
        """空の段落リストでもセクションが作成されることを検証"""
        data = {
            "sections": [
                {
                    "heading": "見出しのみ",
                    "level": 1,
                    "paragraphs": []
                }
            ]
        }
        output_path = os.path.join(temp_output_dir, "empty_paragraphs.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        assert os.path.exists(output_path)

    def test_document_with_title_only(self, word_writer, temp_output_dir):
        """タイトルのみのドキュメントが作成されることを検証"""
        data = {
            "title": "タイトルのみ",
            "sections": []
        }
        output_path = os.path.join(temp_output_dir, "title_only.docx")
        
        result = word_writer.create_document(data, output_path)
        
        assert result.success is True
        assert os.path.exists(output_path)
        
        # タイトルが含まれていることを検証
        doc = Document(output_path)
        title_paragraphs = [p for p in doc.paragraphs if p.style.name == 'Title']
        assert len(title_paragraphs) >= 1
