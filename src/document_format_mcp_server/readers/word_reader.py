"""Word (.docx) file reader."""

import os
import time
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from ..utils.errors import FileNotFoundError, CorruptedFileError
from ..utils.logging_config import get_logger
from ..utils.models import ReadResult, DocumentContent


# ロガーの取得
logger = get_logger(__name__)


class WordReader:
    """Word (.docx)ファイルを読み取るクラス。"""

    def __init__(self, max_file_size_mb: int = 100):
        """
        Wordリーダーを初期化する。

        Args:
            max_file_size_mb: 処理する最大ファイルサイズ（MB）（デフォルト: 100）
        """
        self.max_file_size_mb = max_file_size_mb

    def read_file(self, file_path: str) -> ReadResult:
        """
        Wordファイルからコンテンツを抽出する。

        Args:
            file_path: 読み取るWordファイルのパス

        Returns:
            ReadResult: 読み取り結果を含むデータクラス

        Raises:
            FileNotFoundError: ファイルが存在しない場合
            CorruptedFileError: ファイルが破損している場合
        """
        logger.info(f"Wordファイルの読み込みを開始: {file_path}")
        start_time = time.time()
        
        # ファイルの存在確認
        if not os.path.exists(file_path):
            logger.error(f"ファイルが見つかりません: {file_path}")
            return ReadResult(
                success=False,
                content=None,
                error=f"指定されたファイルが見つかりません: {file_path}",
                file_path=file_path
            )
        
        # ファイルサイズの検証
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        if file_size_mb > self.max_file_size_mb:
            logger.error(
                f"ファイルサイズが制限を超えています: {file_size_mb:.2f}MB > {self.max_file_size_mb}MB"
            )
            return ReadResult(
                success=False,
                content=None,
                error=f"ファイルサイズが制限を超えています: {file_size_mb:.2f}MB（最大: {self.max_file_size_mb}MB）",
                file_path=file_path
            )

        try:
            # Wordファイルを開く
            doc = Document(file_path)
            
            paragraphs_data = []
            tables_data = []
            
            # 段落を処理
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # 空の段落はスキップ
                    para_data = {
                        "text": text,
                        "style": para.style.name,
                        "level": self._get_heading_level(para.style.name)
                    }
                    paragraphs_data.append(para_data)
            
            # 表を処理
            for table in doc.tables:
                table_data = self._extract_table(table)
                tables_data.append(table_data)
            
            content_dict = {
                "paragraphs": paragraphs_data,
                "tables": tables_data
            }
            
            # メタデータを作成
            metadata = {
                "paragraph_count": len(paragraphs_data),
                "table_count": len(tables_data),
                "file_size_mb": file_size_mb
            }
            
            # DocumentContentを作成
            document_content = DocumentContent(
                format_type="docx",
                metadata=metadata,
                content=content_dict
            )
            
            # 処理時間を計算
            elapsed_time = time.time() - start_time
            logger.info(
                f"Wordファイルの読み込みが完了: {file_path} "
                f"(段落数: {len(paragraphs_data)}, 表数: {len(tables_data)}, 処理時間: {elapsed_time:.2f}秒)"
            )
            logger.debug(
                f"抽出されたデータの概要: 段落数={len(paragraphs_data)}, 表数={len(tables_data)}"
            )
            
            # ReadResultを返す
            return ReadResult(
                success=True,
                content=document_content,
                error=None,
                file_path=file_path
            )
        
        except (FileNotFoundError, CorruptedFileError) as e:
            # 既知のエラーはReadResultとして返す
            logger.error(f"Wordファイルの読み込みエラー: {file_path}", exc_info=True)
            return ReadResult(
                success=False,
                content=None,
                error=str(e),
                file_path=file_path
            )
        except PackageNotFoundError as e:
            logger.error(
                f"Wordファイルが破損しています: {file_path}",
                exc_info=True
            )
            return ReadResult(
                success=False,
                content=None,
                error=f"Wordファイルが破損しているか、読み取り不可能です: {file_path}",
                file_path=file_path
            )
        except Exception as e:
            # その他の予期しないエラー
            logger.error(
                f"Wordファイルの読み取り中にエラーが発生: {file_path}",
                exc_info=True
            )
            return ReadResult(
                success=False,
                content=None,
                error=f"Wordファイルの読み取り中にエラーが発生しました: {str(e)}",
                file_path=file_path
            )

    def _get_heading_level(self, style_name: str) -> int:
        """
        スタイル名から見出しレベルを取得する。

        Args:
            style_name: スタイル名（例: "Heading 1", "Heading 2"）

        Returns:
            見出しレベル（1-9）。見出しでない場合は0
        """
        if style_name.startswith("Heading"):
            try:
                # "Heading 1" -> 1, "Heading 2" -> 2, etc.
                level = int(style_name.split()[-1])
                return level
            except (ValueError, IndexError):
                return 0
        return 0

    def _extract_table(self, table) -> dict[str, Any]:
        """
        表からデータを抽出する。

        Args:
            table: python-docxのTableオブジェクト

        Returns:
            表データを含む辞書
        """
        table_data = {
            "rows": len(table.rows),
            "columns": len(table.columns),
            "data": []
        }
        
        # 各行のデータを抽出
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            table_data["data"].append(row_data)
        
        return table_data

