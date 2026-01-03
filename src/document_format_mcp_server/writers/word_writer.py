"""Word (.docx) file writer."""

import os
import re
import time
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..utils.errors import ValidationError
from ..utils.logging_config import get_logger
from ..utils.models import WriteResult


# ロガーの取得
logger = get_logger(__name__)


def _sanitize_text(text: str) -> str:
    """
    XMLに互換性のない制御文字を除去する。
    
    Args:
        text: サニタイズする文字列
        
    Returns:
        制御文字を除去した文字列
    """
    if not isinstance(text, str):
        return str(text) if text is not None else ""
    # XML 1.0で許可されていない制御文字を除去（タブ、改行、キャリッジリターンは許可）
    # 許可: \x09 (tab), \x0A (newline), \x0D (carriage return)
    # 除去: \x00-\x08, \x0B, \x0C, \x0E-\x1F
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F]', '', text)


class WordWriter:
    """Word (.docx)ファイルを生成するクラス。"""

    def create_document(self, data: dict[str, Any], output_path: str) -> WriteResult:
        """
        構造化データからWordファイルを生成する。

        Args:
            data: ドキュメントデータ:
                {
                    "title": str (オプション),
                    "sections": [
                        {
                            "heading": str,
                            "level": int (1-3),
                            "paragraphs": [str, ...],
                            "tables": [
                                {
                                    "data": [[cell_value, ...], ...]
                                }
                            ] (オプション),
                            "bullets": [str, ...] (オプション)
                        }
                    ]
                }
            output_path: 出力ファイルのパス

        Returns:
            作成されたファイルのパス

        Raises:
            ValidationError: 入力データが不正な場合
            Exception: ファイル生成中にエラーが発生した場合
        """
        logger.info(f"Wordファイルの生成を開始: {output_path}")
        start_time = time.time()

        try:
            # 入力データの検証
            self._validate_data(data)

            # 新しいドキュメントを作成
            doc = Document()

            # タイトルを追加（存在する場合）
            if "title" in data and data["title"]:
                # 制御文字をサニタイズ
                sanitized_title = _sanitize_text(data["title"])
                title = doc.add_heading(sanitized_title, level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # セクションを作成
            sections_data = data.get("sections", [])
            for section_data in sections_data:
                self._create_section(doc, section_data)

            # 出力ディレクトリが存在しない場合は作成
            output_dir = os.path.dirname(output_path)
            if output_dir and not os.path.exists(output_dir):
                try:
                    os.makedirs(output_dir, exist_ok=True)
                    logger.debug(f"出力ディレクトリを作成: {output_dir}")
                except OSError as e:
                    logger.error(f"出力ディレクトリの作成に失敗: {output_dir}", exc_info=True)
                    raise Exception(
                        f"出力ディレクトリの作成に失敗しました: {output_dir}"
                    ) from e

            # ファイルを保存
            try:
                doc.save(output_path)
            except PermissionError as e:
                logger.error(f"ファイルの保存に失敗（権限エラー）: {output_path}", exc_info=True)
                raise Exception(
                    f"ファイルの保存に失敗しました（権限エラー）: {output_path}"
                ) from e
            except OSError as e:
                logger.error(f"ファイルの保存に失敗: {output_path}", exc_info=True)
                raise Exception(
                    f"ファイルの保存に失敗しました: {output_path}"
                ) from e

            # 処理時間を計算
            elapsed_time = time.time() - start_time
            logger.info(
                f"Wordファイルの生成が完了: {output_path} "
                f"(セクション数: {len(sections_data)}, 処理時間: {elapsed_time:.2f}秒)"
            )

            # WriteResultを返す
            return WriteResult(
                success=True,
                output_path=output_path,
                url=None,
                error=None
            )

        except ValidationError as e:
            # ValidationErrorはWriteResultとして返す
            logger.error(f"Wordファイルの生成エラー（検証失敗）: {output_path}", exc_info=True)
            return WriteResult(
                success=False,
                output_path=None,
                url=None,
                error=str(e)
            )
        except Exception as e:
            logger.error(
                f"Wordファイルの生成中にエラーが発生: {output_path}",
                exc_info=True
            )
            return WriteResult(
                success=False,
                output_path=None,
                url=None,
                error=f"Wordファイルの生成中にエラーが発生しました: {str(e)}"
            )

    def _validate_data(self, data: dict[str, Any]) -> None:
        """入力データを検証する。"""
        if not isinstance(data, dict):
            raise ValidationError(
                "入力データは辞書形式である必要があります",
                details={"data_type": type(data).__name__}
            )

        if "sections" not in data:
            raise ValidationError(
                "入力データに'sections'キーが必要です",
                details={"keys": list(data.keys())}
            )

        sections = data["sections"]
        if not isinstance(sections, list):
            raise ValidationError(
                "'sections'はリスト形式である必要があります",
                details={"sections_type": type(sections).__name__}
            )

        # 各セクションデータを検証
        for idx, section_data in enumerate(sections):
            if not isinstance(section_data, dict):
                raise ValidationError(
                    f"セクション{idx + 1}のデータは辞書形式である必要があります",
                    details={"section_index": idx, "data_type": type(section_data).__name__}
                )

            # 見出しレベルの検証
            if "level" in section_data:
                level = section_data["level"]
                if not isinstance(level, int) or level < 1 or level > 3:
                    raise ValidationError(
                        f"セクション{idx + 1}の見出しレベルは1-3の整数である必要があります",
                        details={"section_index": idx, "level": level}
                    )

    def _create_section(self, doc: Document, section_data: dict[str, Any]) -> None:
        """セクションを作成する。"""
        # 見出しを追加
        heading = section_data.get("heading", "")
        level = section_data.get("level", 1)
        if heading:
            # 制御文字をサニタイズ
            sanitized_heading = _sanitize_text(heading)
            doc.add_heading(sanitized_heading, level=level)
            logger.debug(f"見出しを追加: {sanitized_heading} (レベル: {level})")

        # 段落を追加
        paragraphs = section_data.get("paragraphs", [])
        for paragraph_text in paragraphs:
            if paragraph_text:
                # 制御文字をサニタイズ
                sanitized_text = _sanitize_text(str(paragraph_text))
                doc.add_paragraph(sanitized_text)

        # 箇条書きリストを追加
        bullets = section_data.get("bullets", [])
        if bullets:
            for bullet_text in bullets:
                if bullet_text:
                    # 制御文字をサニタイズ
                    sanitized_text = _sanitize_text(str(bullet_text))
                    doc.add_paragraph(sanitized_text, style='List Bullet')
            logger.debug(f"箇条書きリストを追加: {len(bullets)}項目")

        # 表を追加
        tables = section_data.get("tables", [])
        for table_data in tables:
            self._create_table(doc, table_data)

    def _create_table(self, doc: Document, table_data: dict[str, Any]) -> None:
        """表を作成する。"""
        data = table_data.get("data", [])
        if not data or not isinstance(data, list):
            logger.warning("表データが空またはリスト形式ではありません")
            return

        # 行数と列数を取得
        rows = len(data)
        if rows == 0:
            return

        cols = len(data[0]) if isinstance(data[0], list) else 1

        # 表を作成
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Light Grid Accent 1'

        # データを設定
        for row_idx, row_data in enumerate(data):
            if not isinstance(row_data, list):
                row_data = [row_data]

            for col_idx, cell_value in enumerate(row_data):
                if col_idx < cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    # 制御文字をサニタイズ
                    sanitized_value = _sanitize_text(str(cell_value) if cell_value is not None else "")
                    cell.text = sanitized_value

        logger.debug(f"表を追加: {rows}行 x {cols}列")
